from __future__ import annotations

import math
import shutil
import tempfile
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageOps


IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".heic",
    ".heif",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}
PDF_EXTENSIONS = {".pdf"}
SUPPORTED_EXTENSIONS = IMAGE_EXTENSIONS | PDF_EXTENSIONS


class ConversionCancelled(Exception):
    pass


@dataclass(frozen=True)
class SourceFile:
    path: Path
    kind: str


@dataclass(frozen=True)
class BuildResult:
    output_path: Path
    source_count: int
    page_count: int


ProgressCallback = Callable[[int, int, str], None]


def scan_sources(root: Path, max_depth: int | None) -> list[SourceFile]:
    root = root.expanduser().resolve()
    if not root.is_dir():
        raise FileNotFoundError(str(root))

    sources: list[SourceFile] = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            continue
        depth = len(path.relative_to(root).parent.parts)
        if max_depth is not None and depth > max_depth:
            continue
        kind = "pdf" if suffix in PDF_EXTENSIONS else "image"
        sources.append(SourceFile(path=path, kind=kind))

    return sorted(sources, key=lambda item: str(item.path.relative_to(root)).lower())


def build_word(
    root: Path,
    output_path: Path,
    max_depth: int | None,
    progress: ProgressCallback | None = None,
    cancel_event: threading.Event | None = None,
) -> BuildResult:
    root = root.expanduser().resolve()
    output_path = output_path.expanduser().resolve()
    sources = scan_sources(root, max_depth)
    if not sources:
        raise ValueError("no_sources")

    tmp_dir = Path(tempfile.mkdtemp(prefix="image_pdf_word_"))
    rendered: list[Path] = []

    try:
        total = len(sources)
        for index, source in enumerate(sources, start=1):
            _raise_if_cancelled(cancel_event)
            if progress:
                progress(index - 1, total, source.path.name)
            if source.kind == "image":
                rendered.append(_normalize_image(source.path, tmp_dir / f"{index:04d}.jpg"))
            else:
                rendered.extend(_render_pdf(source.path, tmp_dir, index, cancel_event))

        _raise_if_cancelled(cancel_event)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        _create_docx(rendered, output_path)
        if progress:
            progress(total, total, output_path.name)
        return BuildResult(output_path=output_path, source_count=len(sources), page_count=len(rendered))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _raise_if_cancelled(cancel_event: threading.Event | None) -> None:
    if cancel_event and cancel_event.is_set():
        raise ConversionCancelled()


def _normalize_image(src: Path, dest: Path) -> Path:
    with Image.open(src) as image:
        image = ImageOps.exif_transpose(image)
        if image.mode in ("RGBA", "LA") or "transparency" in image.info:
            rgba = image.convert("RGBA")
            background = Image.new("RGB", rgba.size, "white")
            background.paste(rgba, mask=rgba.split()[-1])
            image = background
        else:
            image = image.convert("RGB")
        image.save(dest, "JPEG", quality=94, optimize=True)
    return dest


def _render_pdf(
    src: Path,
    tmp_dir: Path,
    source_index: int,
    cancel_event: threading.Event | None,
) -> list[Path]:
    rendered: list[Path] = []
    document = fitz.open(src)
    try:
        for page_index in range(document.page_count):
            _raise_if_cancelled(cancel_event)
            page = document.load_page(page_index)
            matrix = fitz.Matrix(2.5, 2.5)
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            png_path = tmp_dir / f"{source_index:04d}_{page_index + 1:03d}.png"
            jpg_path = tmp_dir / f"{source_index:04d}_{page_index + 1:03d}.jpg"
            pixmap.save(png_path)
            _normalize_image(png_path, jpg_path)
            png_path.unlink(missing_ok=True)
            rendered.append(jpg_path)
    finally:
        document.close()
    return rendered


def _create_docx(images: Iterable[Path], output_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    for index, image_path in enumerate(images):
        _add_image_page(doc, image_path, first=(index == 0))

    doc.save(output_path)


def _add_image_page(doc: Document, image_path: Path, first: bool) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    landscape = width_px > height_px * 1.08
    section = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section(section, landscape)

    max_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    max_height = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
    scale = min(max_width / width_px, max_height / height_px)
    picture_width = max(0.1, math.floor(width_px * scale * 1000 - 3) / 1000)
    picture_height = max(0.1, math.floor(height_px * scale * 1000 - 3) / 1000)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(
        str(image_path),
        width=Inches(picture_width),
        height=Inches(picture_height),
    )


def _set_section(section, landscape: bool) -> None:
    section.page_width = Inches(11.69 if landscape else 8.27)
    section.page_height = Inches(8.27 if landscape else 11.69)
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.22)
    section.bottom_margin = Inches(0.22)
    section.left_margin = Inches(0.22)
    section.right_margin = Inches(0.22)
